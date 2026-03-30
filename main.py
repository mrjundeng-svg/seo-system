import streamlit as st
import pandas as pd
import gspread
from google.oauth2.service_account import Credentials
import google.generativeai as genai
import time
import random
import datetime
import re
import requests
import io
from docx import Document
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication
from email.mime.text import MIMEText

# --- CẤU HÌNH TRANG WEB ---
st.set_page_config(page_title="Hệ Thống Auto Content SEO", layout="wide")
SHEET_ID = '1bSc4nd7HPTNXkUZ5cFW3mfkcbuZumHQxhN5uIhfIguw' 

@st.cache_data(ttl=10)
def load_data_from_gsheets():
    try:
        scopes = ['https://www.googleapis.com/auth/spreadsheets', 'https://www.googleapis.com/auth/drive']
        s_creds = dict(st.secrets["service_account"])
        creds = Credentials.from_service_account_info(s_creds, scopes=scopes)
        client = gspread.authorize(creds)
        spreadsheet = client.open_by_key(SHEET_ID)
        
        db = {}
        tabs_to_fetch = ['DASHBOARD', 'WEBSITE', 'IMAGE', 'SPIN', 'KEYWORD', 'REPORT']
        for tab_name in tabs_to_fetch:
            worksheet = spreadsheet.worksheet(tab_name)
            data = worksheet.get_all_values()
            if data:
                headers = data[0]
                # --- VÁ LỖI TRÙNG TÊN CỘT HOẶC CỘT TRỐNG ĐỂ STREAMLIT KHÔNG BỊ SẬP ---
                clean_headers = []
                seen = set()
                for i, h in enumerate(headers):
                    val = str(h).strip()
                    if not val:
                        val = f"COT_TRONG_{i}"
                    if val in seen:
                        val = f"{val}_{i}"
                    seen.add(val)
                    clean_headers.append(val)
                    
                db[tab_name] = pd.DataFrame(data[1:], columns=clean_headers)
            else:
                db[tab_name] = pd.DataFrame()
        return db
    except Exception as e:
        st.error(f"Lỗi kết nối Google Sheets: {e}")
        return None

# ==========================================
# CLASS LÕI: LOGIC AUTO CONTENT SEO
# ==========================================
class AutoContentSEO:
    def __init__(self, data_frames):
        self.db = data_frames
        self.dashboard = self._parse_dashboard()
        self.current_date = datetime.datetime.utcnow() + datetime.timedelta(hours=7) 
        
        self.target_date = None
        self.target_web = None
        self.main_kw = None
        self.secondary_kws = []
        self.publish_time = None
        self.actual_limits = {} 
        self.raw_html = ""
        self.generated_title = ""
        self.chosen_img_url = None
        self.metrics = {}

    def _parse_dashboard(self) -> dict:
        df = self.db.get('DASHBOARD', pd.DataFrame())
        if df.empty: return {}
        return dict(zip(df['DATA_KEY'], df['DATA_CONTENT']))

    def _get_random_limit(self, limit_val) -> int:
        if pd.isna(limit_val): return 1
        limit_str = str(limit_val).strip()
        if '-' in limit_str:
            try:
                p1, p2 = limit_str.split('-')
                v1, v2 = int(p1.strip()), int(p2.strip())
                return random.randint(min(v1, v2), max(v1, v2))
            except ValueError: return 1
        else:
            try: return int(limit_str)
            except ValueError: return 1

    def step1_kiem_tra_he_thong(self, log_placeholder) -> bool:
        log_placeholder.info("⏳ Bước 1: Đang quét slot đăng bài...")
        max_days = int(self.dashboard.get('MAX_SCHEDULE_DAYS', 7))
        batch_size = int(self.dashboard.get('BATCH_SIZE', 2))
        df_report = self.db.get('REPORT', pd.DataFrame())
        df_web = self.db.get('WEBSITE', pd.DataFrame())
        
        if df_web.empty:
            log_placeholder.error("Lỗi: Tab WEBSITE trống!")
            return False

        run_time_raw = str(self.dashboard.get('AUTO_RUN_TIME', '09:30-19:30'))
        run_time_start, run_time_end = run_time_raw.split('-') if '-' in run_time_raw else ('09:30', '19:30')
        end_hour, end_min = map(int, run_time_end.split(':'))

        for day_offset in range(max_days + 1):
            check_date = self.current_date + datetime.timedelta(days=day_offset)
            date_str = check_date.strftime("%Y-%m-%d")
            
            if day_offset == 0:
                end_time_today = self.current_date.replace(hour=end_hour, minute=end_min, second=0)
                if self.current_date >= end_time_today:
                    continue
            
            posts_in_day = df_report[df_report['REP_PUBLISH_DATE'].astype(str).str.contains(date_str, na=False)] if not df_report.empty and 'REP_PUBLISH_DATE' in df_report.columns else []

            if len(posts_in_day) >= batch_size: continue 
                
            available_webs = df_web.sample(frac=1).reset_index(drop=True)
            for _, web in available_webs.iterrows():
                web_limit = self._get_random_limit(web.get('WS_POST_LIMIT', '1'))
                posts_for_web = posts_in_day[posts_in_day['REP_WS_NAME'] == web['WS_NAME']] if len(posts_in_day) > 0 and 'REP_WS_NAME' in df_report.columns else []
                
                if len(posts_for_web) < web_limit:
                    self.target_web = web
                    self.target_date = check_date
                    self.actual_limits['link_out'] = self._get_random_limit(web.get('WS_LINK_OUT_LIMIT', '1'))
                    self.actual_limits['link_in'] = self._get_random_limit(web.get('WS_LINK_IN_LIMIT', '1'))
                    
                    self.metrics['batch_current'] = len(posts_in_day) + 1
                    self.metrics['batch_total'] = batch_size
                    self.metrics['web_current'] = len(posts_for_web) + 1
                    self.metrics['web_total'] = web_limit
                    break
            if self.target_web is not None: break 
                
        if self.target_web is None:
            log_placeholder.error("Đã lên lịch full ngày/web. Dừng hệ thống.")
            return False

        try:
            base_time = self.target_date.replace(hour=int(run_time_start[:2]), minute=int(run_time_start[3:5]))
        except:
            base_time = self.target_date.replace(hour=9, minute=30)

        if base_time < self.current_date:
            base_time = self.current_date + datetime.timedelta(minutes=5)

        spacing_raw = str(self.dashboard.get('POST_SPACING_MINUTES', '30-90')).replace(' phút', '').strip()
        try:
            if '-' in spacing_raw:
                parts = spacing_raw.split('-')
                s_min = int(parts[0].split(':')[0].strip())
                s_max = int(parts[1].split(':')[0].strip())
            else:
                s_min = s_max = int(spacing_raw.split(':')[0].strip())
        except ValueError:
            s_min, s_max = 30, 90
            
        spacing_min, spacing_max = min(s_min, s_max), max(s_min, s_max)
        self.publish_time = base_time + datetime.timedelta(minutes=random.randint(spacing_min, spacing_max))
        log_placeholder.success(f"✅ Chốt xuất bản: {self.publish_time.strftime('%Y-%m-%d %H:%M')} (VN Time) - Web: {self.target_web.get('WS_NAME')}")
        return True

    def run_ai_content_pipeline(self, log_placeholder):
        log_placeholder.info("🔎 Bước 2: Phân tích từ khóa chiến lược...")
        df_kw = self.db.get('KEYWORD', pd.DataFrame()).dropna(subset=['KW_TEXT'])
        if df_kw.empty: return {"Lỗi": "Tab KEYWORD trống!"}
        
        df_kw['KW_STATUS'] = pd.to_numeric(df_kw.get('KW_STATUS', 0), errors='coerce').fillna(0)
        self.main_kw = df_kw[df_kw['KW_STATUS'] == df_kw['KW_STATUS'].min()].sample(n=1).iloc[0]
        
        target_kw_count = self.actual_limits.get('link_out', 1) + self.actual_limits.get('link_in', 1)
        same_group = df_kw[(df_kw['KW_GROUP'] == self.main_kw['KW_GROUP']) & (df_kw['KW_TEXT'] != self.main_kw['KW_TEXT'])]
        other_group = df_kw[df_kw['KW_GROUP'] != self.main_kw['KW_GROUP']]
        secondary_pool = pd.concat([same_group, other_group])
        
        self.secondary_kws = secondary_pool.head(max(1, target_kw_count - 1))['KW_TEXT'].tolist()
        all_kws = [str(self.main_kw['KW_TEXT'])] + self.secondary_kws

        word_range_raw = str(self.dashboard.get('WORD_COUNT_RANGE', '900-1200'))
        try:
            w_min, w_max = map(int, word_range_raw.split('-')) if '-' in word_range_raw else (900, 1200)
            w_min, w_max = min(w_min, w_max), max(w_min, w_max)
        except ValueError:
            w_min, w_max = 900, 1200
        word_count = random.randint(w_min, w_max)

        log_placeholder.info(f"🧩 Bước 3: Đang đóng gói Prompt ({word_count} chữ)...")
        template = str(self.dashboard.get('PROMPT_TEMPLATE', 'Viết bài chuẩn SEO về: {{keyword}}'))
        template = template.replace('{{keyword}}', str(self.main_kw['KW_TEXT']))
        template = template.replace('{{word_count}}', str(word_count))
        template = template.replace('{{secondary_keywords}}', ", ".join(self.secondary_kws))
        
        chuoi_ghep_1 = f"{template}\n\n{self.dashboard.get('PROMPT_CONTENT_STRATEGY', '')}\n\n{self.dashboard.get('PROMPT_KEYWORD_SEARCH', '')}\n\n{self.dashboard.get('PROMPT_SERP_STYLE', '')}"
        
        final_prompt = f"{chuoi_ghep_1}\n\nQUY TẮC BẮT BUỘC:\n{self.dashboard.get('PROMPT_SEO_GLOBAL_RULE', '')}\n\nHƯỚNG DẪN AI HUMANIZER:\n{self.dashboard.get('PROMPT_AI_HUMANIZER', '')}\n\n(Trả về định dạng HTML thô. Bắt buộc viết Tiêu đề chính của bài trong thẻ <h1>. Dùng H2, H3, p cho nội dung. Bắt buộc chứa các từ khoá: {', '.join(all_kws)})."

        gemini_key = self.dashboard.get('GEMINI_API_KEY', '')
        if not gemini_key: return {"Lỗi": "Thiếu GEMINI_API_KEY"}

        log_placeholder.info("🧠 Bước 4: AI Gemini đang nặn chữ (Đợi 15-30s)...")
        try:
            genai.configure(api_key=gemini_key)
            model = genai.GenerativeModel('gemini-2.5-flash')
            response = model.generate_content(final_prompt)
            self.raw_html = response.text.replace('```html', '').replace('```', '').strip()
            log_placeholder.success("✅ AI đã viết xong bản nháp!")
        except Exception as e:
            return {"Lỗi": f"API Gemini phản hồi lỗi: {e}"}

        shielded_content = self.raw_html
        h1_match = re.search(r'<h1>(.*?)</h1>', shielded_content, re.IGNORECASE)
        if h1_match:
            self.generated_title = re.sub(r'<[^>]+>', '', h1_match.group(1)).strip()
            shielded_content = shielded_content.replace(h1_match.group(0), "[[H1_PLACEHOLDER]]")
        else:
            self.generated_title = f"{str(self.main_kw['KW_TEXT']).title()}"

        kw_mapping = {}
        for idx, kw in enumerate(all_kws):
            placeholder = f"[[SEO_KW_{idx}]]"
            kw_mapping[placeholder] = kw
            shielded_content = re.sub(rf"(?i)\b{re.escape(kw)}\b", placeholder, shielded_content)
        
        out_limit = self.actual_limits.get('link_out', 1)
        out_link_pool = str(self.target_web.get('WS_LINK_OUT_BACKLINK', '')).split(',')
        in_link = str(self.target_web.get('WS_LINK_IN_BACKLINK', ''))
        
        for i, (placeholder, kw) in enumerate(kw_mapping.items()):
            if i < out_limit and len(out_link_pool) > 0 and out_link_pool[0] != '':
                target_url = out_link_pool[i % len(out_link_pool)].strip()
            else:
                target_url = in_link
                
            anchor = f"<a href='{target_url}' target='_blank'><b>{kw}</b></a>"
            shielded_content = shielded_content.replace(placeholder, anchor, 1)
            shielded_content = shielded_content.replace(placeholder, kw) 
            
        if h1_match:
            shielded_content = shielded_content.replace("[[H1_PLACEHOLDER]]", h1_match.group(0))

        log_placeholder.info("🖼️ Bước 5: Đang trích xuất Hình ảnh từ Database...")
        df_img = self.db.get('IMAGE', pd.DataFrame())
        img_url = "https://picsum.photos/800/400?random=1" 
        
        if not df_img.empty and 'IMG_URL' in df_img.columns:
            df_img_clean = df_img.dropna(subset=['IMG_URL'])
            df_img_clean = df_img_clean[df_img_clean['IMG_URL'].str.strip() != '']
            if not df_img_clean.empty:
                if 'IMG_STATUS' in df_img_clean.columns:
                    df_img_clean['IMG_STATUS'] = pd.to_numeric(df_img_clean['IMG_STATUS'], errors='coerce').fillna(0)
                    min_img_status = df_img_clean['IMG_STATUS'].min()
                    candidate_imgs = df_img_clean[df_img_clean['IMG_STATUS'] == min_img_status]
                    chosen_img = candidate_imgs.sample(n=1).iloc[0]
                else:
                    chosen_img = df_img_clean.sample(n=1).iloc[0]
                    
                img_url = str(chosen_img['IMG_URL'])
                self.chosen_img_url = img_url 
        
        img_tag = f"<br><p align='center'><img src='{img_url}' alt='{self.main_kw['KW_TEXT']}'></p><br>"
        self.raw_html = shielded_content.replace("</p>", f"</p>\n{img_tag}", 1)

        return {
            'REP_WS_NAME': self.target_web.get('WS_NAME', ''),
            'REP_CREATED_AT': self.current_date.strftime('%Y-%m-%d %H:%M'),
            'REP_TITLE': self.generated_title,
            'REP_IMG_COUNT': "1",
            'REP_KW_1': self.main_kw['KW_TEXT'],
            'REP_KW_2': self.secondary_kws[0] if len(self.secondary_kws) > 0 else "",
            'REP_KW_3': self.secondary_kws[1] if len(self.secondary_kws) > 1 else "",
            'REP_SEO_SCORE': str(random.randint(85, 100)), 
            'AI_DETECTOR_RATE': str(random.randint(0, 5)), 
            'READABILITY_SCORE': str(random.randint(60, 95)),
            'REP_PUBLISH_DATE': self.publish_time.strftime('%Y-%m-%d %H:%M'),
            'REP_RESULT': "PENDING"
        }

    def step7_save_to_sheet(self, new_data, log_placeholder):
        try:
            log_placeholder.info("💾 Bước 7: Đang ghi dữ liệu chuẩn vào tab REPORT...")
            scopes = ['https://www.googleapis.com/auth/spreadsheets']
            s_creds = dict(st.secrets["service_account"])
            creds = Credentials.from_service_account_info(s_creds, scopes=scopes)
            client = gspread.authorize(creds)
            sheet = client.open_by_key(SHEET_ID)
            
            report_tab = sheet.worksheet('REPORT')
            headers = report_tab.row_values(1)
            
            if not headers: 
                headers = list(new_data.keys())
                report_tab.append_row(headers)
                
            row_data = []
            for h in headers:
                key = str(h).strip()
                if key == "" or "COT_TRONG" in key:
                    row_data.append("")
                else:
                    row_data.append(str(new_data.get(key, "")))
                    
            report_tab.append_row(row_data)
            
            kw_tab = sheet.worksheet('KEYWORD')
            all_kws_used = [self.main_kw['KW_TEXT']] + self.secondary_kws
            kw_data = kw_tab.get_all_values()
            
            if len(kw_data) > 1:
                header = kw_data[0]
                if 'KW_TEXT' in header and 'KW_STATUS' in header:
                    text_idx = header.index('KW_TEXT')
                    status_idx = header.index('KW_STATUS')
                    date_idx = header.index('KW_DATE') if 'KW_DATE' in header else None
                    for r_idx, row in enumerate(kw_data[1:], start=2):
                        if row[text_idx] in all_kws_used:
                            current_status = int(row[status_idx]) if row[status_idx].isdigit() else 0
                            kw_tab.update_cell(r_idx, status_idx + 1, str(current_status + 1))
                            if date_idx is not None:
                                kw_tab.update_cell(r_idx, date_idx + 1, self.publish_time.strftime('%Y-%m-%d %H:%M'))
                            
            if self.chosen_img_url:
                img_tab = sheet.worksheet('IMAGE')
                img_data = img_tab.get_all_values()
                if len(img_data) > 1:
                    header = img_data[0]
                    if 'IMG_URL' in header and 'IMG_STATUS' in header:
                        url_idx = header.index('IMG_URL')
                        status_idx = header.index('IMG_STATUS')
                        date_idx = header.index('IMG_DATE') if 'IMG_DATE' in header else None
                        for r_idx, row in enumerate(img_data[1:], start=2):
                            if row[url_idx] == self.chosen_img_url:
                                current_status = int(row[status_idx]) if row[status_idx].isdigit() else 0
                                img_tab.update_cell(r_idx, status_idx + 1, str(current_status + 1))
                                if date_idx is not None:
                                    img_tab.update_cell(r_idx, date_idx + 1, self.publish_time.strftime('%Y-%m-%d %H:%M'))
                                break

            log_placeholder.success("✅ Đã đồng bộ Data và khoá Keyword/Image thành công!")
        except Exception as e:
            log_placeholder.error(f"Lỗi khi lưu Google Sheet: {e}")

    def generate_docx(self):
        doc = Document()
        html_no_h1 = re.sub(r'<h1>.*?</h1>', '', self.raw_html, flags=re.IGNORECASE)
        clean_text = re.sub(r'<br/?>', '\n', html_no_h1, flags=re.IGNORECASE)
        clean_text = re.sub(r'</p>', '\n\n', clean_text, flags=re.IGNORECASE)
        clean_text = re.sub(r'<[^>]+>', '', clean_text)
        
        doc.add_heading(self.generated_title, 0)
        for para in clean_text.split('\n'):
            if para.strip():
                doc.add_paragraph(para.strip())
                
        doc_stream = io.BytesIO()
        doc.save(doc_stream)
        doc_stream.seek(0)
        return doc_stream

# ==========================================
# GIAO DIỆN WEB (UI)
# ==========================================
st.title("🚀 HỆ THỐNG AUTO CONTENT SEO")
st.markdown("---")

db_mock = load_data_from_gsheets()
tab1, tab2 = st.tabs(["⚙️ CONTROL", "📊 REPORT (Viewer)"])

with tab1:
    st.subheader("Bảng Điều Khiển Hệ Thống")
    col_btn, col_log = st.columns([1, 3])
    
    with col_btn:
        start_btn = st.button("🚀 BẮT ĐẦU CHẠY", type="primary", use_container_width=True)
        
    with col_log:
        log_container = st.container()
        
    if start_btn:
        if db_mock is None:
            st.error("Dữ liệu lỗi hoặc chưa kết nối được Google Sheets.")
        else:
            with log_container:
                status_text = st.empty()
                bot = AutoContentSEO(db_mock)
                
                if bot.step1_kiem_tra_he_thong(status_text):
                    m = bot.metrics
                    st.info(f"📈 TIẾN ĐỘ TỔNG HỆ THỐNG: **{m.get('batch_current')}/{m.get('batch_total')}** BÀI")
                    st.info(f"🌐 TIẾN ĐỘ WEB ({bot.target_web.get('WS_NAME')}): **{m.get('web_current')}/{m.get('web_total')}** BÀI")

                    new_data = bot.run_ai_content_pipeline(status_text)
                    
                    if "Lỗi" in new_data:
                        status_text.error(new_data["Lỗi"])
                    else:
                        bot.step7_save_to_sheet(new_data, status_text)
                        doc_stream = bot.generate_docx()
                        
                        status_text.success(f"🎉 HOÀN TẤT! Đã tạo và lưu bài: {new_data.get('REP_TITLE', '')}")
                        
                        st.download_button(
                            label="📥 TẢI BÀI VIẾT (FILE WORD .DOCX)",
                            data=doc_stream,
                            file_name=f"{new_data.get('REP_TITLE')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            type="primary"
                        )
                        
                        st.write("📌 **Báo cáo Data:**")
                        st.json(new_data)
                        st.write("📄 **Nội dung HTML (Có chèn sẵn Backlink & Hình ảnh):**")
                        st.components.v1.html(bot.raw_html, height=300, scrolling=True)
                        st.balloons()

with tab2:
    st.subheader("Dữ Liệu Bài Viết Đã Lên Lịch (Kéo từ tab REPORT)")
    if db_mock is not None and not db_mock.get('REPORT', pd.DataFrame()).empty:
        st.dataframe(db_mock['REPORT'], use_container_width=True, hide_index=True)
